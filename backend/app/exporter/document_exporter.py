from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    qn = None
    Inches = None
    Pt = None
    RGBColor = None


class DocumentExporter:
    """Exports the current writing task to txt or docx."""

    def __init__(self, export_dir: Path):
        self.export_dir = export_dir
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def export(self, context: dict[str, Any], file_format: str = "docx") -> dict[str, Any]:
        file_format = (file_format or "docx").lower().strip(".")
        if file_format == "txt":
            path = self._export_txt(context)
        elif file_format == "docx":
            path = self._export_docx(context)
        else:
            raise ValueError("Unsupported export format. Use docx or txt.")
        return {
            "format": file_format,
            "fileName": path.name,
            "path": str(path.resolve()),
            "sizeBytes": path.stat().st_size,
        }

    def _safe_name(self, context: dict[str, Any], suffix: str) -> Path:
        task_id = context.get("taskId", "task")
        title = context.get("outline", {}).get("title") or context.get(
            "requirement", {}
        ).get("topic", "document")
        safe_title = "".join(
            ch for ch in str(title)[:30] if ch not in '\\/:*?"<>|\r\n'
        ).strip()
        return self.export_dir / f"{task_id}_{safe_title or 'document'}.{suffix}"

    def _export_txt(self, context: dict[str, Any]) -> Path:
        path = self._safe_name(context, "txt")
        path.write_text(self._plain_content(context), encoding="utf-8")
        return path

    def _export_docx(self, context: dict[str, Any]) -> Path:
        if Document is None:
            return self._export_txt(context)

        path = self._safe_name(context, "docx")
        document = Document()
        self._apply_document_style(document)

        title = context.get("outline", {}).get("title") or context.get(
            "requirement", {}
        ).get("topic", "文档")
        self._add_title(document, str(title))

        requirement = context.get("requirement", {})
        meta = document.add_paragraph(style="Subtitle")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(
            f"文体：{requirement.get('genre', '未设置')}  "
            f"目标字数：{requirement.get('wordCount', '未设置')}  "
            f"风格：{requirement.get('style', '未设置')}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(89, 89, 89)

        outline = context.get("outline", {})
        if outline:
            self._add_outline(document, outline)

        document.add_heading("正文", level=1)
        draft = context.get("draft", {}).get("content", "")
        for paragraph in self._draft_paragraphs(draft, str(title)):
            document.add_paragraph(paragraph)

        summary = context.get("summary")
        if summary:
            document.add_heading("摘要", level=2)
            document.add_paragraph(summary)

        footer = document.sections[0].footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("文档智能写作助手生成")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(127, 127, 127)

        document.save(path)
        return path

    def _plain_content(self, context: dict[str, Any]) -> str:
        title = context.get("outline", {}).get("title") or context.get(
            "requirement", {}
        ).get("topic", "文档")
        outline = context.get("outline", {})
        draft = context.get("draft", {}).get("content", "")
        summary = context.get("summary", "")
        parts = [str(title), ""]
        if outline:
            parts.append("写作提纲")
            if outline.get("thesis"):
                parts.append(f"中心思想：{outline['thesis']}")
            for section in outline.get("sections", []):
                parts.append(section.get("heading", "段落"))
                for point in section.get("points", []):
                    parts.append(f"- {point}")
            parts.append("")
        parts.extend(["正文", *self._draft_paragraphs(draft, str(title))])
        if summary:
            parts.extend(["", "摘要", summary])
        return "\n".join(parts)

    def _apply_document_style(self, document: Any) -> None:
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        self._set_east_asia_font(normal, "Microsoft YaHei")
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        for style_name, size, color in [
            ("Heading 1", 16, RGBColor(46, 116, 181)),
            ("Heading 2", 13, RGBColor(46, 116, 181)),
            ("Heading 3", 12, RGBColor(31, 77, 120)),
        ]:
            style = styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = color
            self._set_east_asia_font(style, "Microsoft YaHei")
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(4)

    def _add_title(self, document: Any, title: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 58, 95)
        self._set_east_asia_run_font(run, "Microsoft YaHei")

    def _add_outline(self, document: Any, outline: dict[str, Any]) -> None:
        document.add_heading("写作提纲", level=1)
        thesis = outline.get("thesis")
        if thesis:
            paragraph = document.add_paragraph()
            paragraph.add_run("中心思想：").bold = True
            paragraph.add_run(str(thesis))

        for section in outline.get("sections", []):
            heading = section.get("heading", "段落")
            document.add_paragraph(str(heading), style="List Number")
            for point in section.get("points", []):
                document.add_paragraph(str(point), style="List Bullet")

    def _draft_paragraphs(self, draft: str, title: str) -> list[str]:
        paragraphs = [p.strip() for p in draft.split("\n\n") if p.strip()]
        if paragraphs and paragraphs[0].strip() == title.strip():
            paragraphs = paragraphs[1:]
        return paragraphs

    def _set_east_asia_font(self, style: Any, font_name: str) -> None:
        if qn is None:
            return
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr._add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)

    def _set_east_asia_run_font(self, run: Any, font_name: str) -> None:
        if qn is None:
            return
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = rpr._add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
